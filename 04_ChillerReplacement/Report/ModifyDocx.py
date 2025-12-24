# _*_ coding: utf-8 _*_
"""
Time:     6/8/2023 12:30 PM
Author:   XuLing
"""
import os
import time

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.table import _Cell
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from win32com.client import Dispatch

class ModifyDocx:
    def __init__(self, GetData):
        self.GetData = GetData
        self.doc_file = docx.Document(self.GetData.docx_path)

    def modify_cover(self, replace_cover):
        """遍历全文的文本框，并修改封面中的文本框里的project名字
        Returns
        -------
        """
        for key, value in replace_cover.items():
            for i in self.doc_file.element.body.iter():
                if i.tag.endswith('txbx'): #txbx查找全文的文本框
                    for j in i.iter():
                        if j.tag.endswith('main}r'):
                            if j.text == key:
                                j.text = value

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def add_table(self, rows, cols, style="table style1"):
        table = self.doc_file.add_table(rows, cols, style=style)
        return table

    def set_cell_font(self, table, pt=8):
        for row in table.rows:
            for cell in row.cells:
                paras = cell.paragraphs
                for para in paras:
                    for run in para.runs:
                        font = run.font
                        font.size = docx.shared.Pt(pt)

    def replace_cell(self, table, change_rows, change_cols, target_values):
        """
        :param table: 需要修改的table
        :param change_rows: 需要修改的table的行
        :param change_cols: 需要修改的table的列
        :param target_values: 需填充的目标值
        :return:
        """
        k = 0
        for i in change_rows:
            for j in change_cols:
                cell = table.cell(i, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #表格内容垂直居中
                para = cell.paragraphs[0] #如果直接使用cell.text来填充，会自带缩进
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER  ##水平居中
                para.text = str(target_values[k])
                run_obj = para.runs[0]
                font = run_obj.font
                font.size = Pt(10)
                k += 1

    def change_imgs(self, add_pics_path):
        for para in self.doc_file.paragraphs:
            for key, pic_paths in add_pics_path.items():
                if key in para.text:
                    # print(key)
                    for pic_path in pic_paths:
                        run = para.add_run('')
                        run.add_picture(pic_path, width=docx.shared.Inches(6.3))

    def replace_placeholder(self, replace_text):
        """替换占位符"""
        self.doc_file.save(os.path.abspath(self.GetData.save_path))
        app = Dispatch('Word.Application')
        doc = app.Documents.Open(os.path.abspath(self.GetData.save_path))

        for paragraph in doc.paragraphs:
            for key, value in replace_text.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            run.italic = False
        doc.Save()



    def replace_text(self, replace_text, del_section):
        self.doc_file.save(os.path.abspath(self.GetData.save_path))

        import pythoncom
        pythoncom.CoInitialize()
        app = Dispatch('Word.Application')# 加上的
        pythoncom.CoInitialize()  # 加上的

        doc = app.Documents.Open(os.path.abspath(self.GetData.save_path))
        s = app.Selection
        if isinstance(del_section, list):
            for key in del_section:
                s.Find.Execute(key, False, False, True, False, False, True, 0, False, "", 2)  #删除一段文字
        for key, value in replace_text.items():
            # print("relace_text", key, value)
            s.Find.Execute(key, False, False, False, False, False, True, 0, False, value, 2) #根据dict的key和value来替代文字
        # time.sleep(2)
        # doc.TablesOfContents(1).Update()
        doc.Save()
        doc.Close()

    def del_tables(self, del_index):
        if self.GetData.chillertype[0:2] == "19":
            type = "centrifugal"
        elif not os.access(self.GetData.data_B_path, os.F_OK):
            type = "single_loop_screw"
        else:
            type = "screw"
        for index, values in del_index.items():
            for value in values:
                if type == index:
                    del_table = self.doc_file.tables[value]
                    del_table._element.getparent().remove(del_table._element)

    def del_blank_pages(self): #此函数用于删除docx文档的空白页
        for paragraph in self.doc_file.paragraphs:
            #如果段落为空且没有任何文本，则删除该段落
            if len(paragraph.text.strip()) == 0 and len(paragraph.runs) == 0:
                self.doc_file.element.body.remove(paragraph._element)


